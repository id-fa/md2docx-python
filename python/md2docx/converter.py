"""IR → Word ドキュメント変換器。

Rust 版 `src/converter.rs` の Python 移植。
- 見出しカウンタ管理 (heading.HeadingManager)
- 図表番号の生成 (sequential = SEQ フィールド / chapter = プレーンテキスト)
- 表幅の動的計算とセル余白
- 画像の本文幅縮小 (Pillow で PNG 化 + 寸法計算)
- 英日間スペース削除
- ハイパーリンク (外部 URL / アンカー両対応)
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObj
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu
from lxml import etree  # type: ignore[import-not-found]
from PIL import Image as PILImage

from . import ir
from . import styles
from .config import Config, PageConfig
from .heading import HeadingManager


EMU_PER_PIXEL = 9_525
EMU_PER_TWIP = 635
TABLE_WIDTH_PCT = 5_000
TABLE_CELL_PADDING_TWIP = 80


class _InlineStyle(Enum):
    BODY = "body"
    TABLE_BODY = "table_body"
    TABLE_HEADER = "table_header"


@dataclass
class _RunFontSpec:
    ascii_font: str
    east_asia: str
    size_pt: float


def convert_to_docx(
    blocks: list[ir.Block], config: Config, base_path: Path
) -> DocumentObj:
    """IR を Word ドキュメントに変換する。"""
    document = Document()
    styles.setup_document_styles(document, config)
    _apply_page_settings(document, config.page)

    ctx = _ConvertContext(config=config, base_path=base_path, document=document)
    for block in blocks:
        ctx.convert_block(block)
    return document


def _apply_page_settings(document: DocumentObj, page: PageConfig) -> None:
    """ページサイズと余白を section に設定する。"""
    section = document.sections[0]
    section.page_width = Emu(_twip_to_emu(page.width))
    section.page_height = Emu(_twip_to_emu(page.height))
    section.top_margin = Emu(_twip_to_emu(page.margin_top))
    section.right_margin = Emu(_twip_to_emu(page.margin_right))
    section.bottom_margin = Emu(_twip_to_emu(page.margin_bottom))
    section.left_margin = Emu(_twip_to_emu(page.margin_left))
    section.header_distance = Emu(_twip_to_emu(page.margin_header))
    section.footer_distance = Emu(_twip_to_emu(page.margin_footer))
    section.gutter = Emu(_twip_to_emu(max(page.margin_gutter, 0)))


def _twip_to_emu(twip: int) -> int:
    return int(twip) * EMU_PER_TWIP


# --------- 変換コンテキスト ---------

class _ConvertContext:
    def __init__(self, *, config: Config, base_path: Path, document: DocumentObj) -> None:
        self.config = config
        self.base_path = base_path
        self.document = document
        self.heading_mgr = HeadingManager()
        self.chapter_number = 0
        self.figure_in_chapter = 0
        self.table_in_chapter = 0
        self.figure_seq = 0
        self.table_seq = 0

    # ---- ディスパッチ ----

    def convert_block(self, block: ir.Block) -> None:
        if isinstance(block, ir.Heading):
            self._convert_heading(block.level, block.content)
        elif isinstance(block, ir.PageBreak):
            self._convert_page_break()
        elif isinstance(block, ir.Paragraph):
            self._convert_paragraph(block.content)
        elif isinstance(block, ir.BulletList):
            self._convert_bullet_list(block.items, depth=0)
        elif isinstance(block, ir.OrderedList):
            self._convert_ordered_list(block.items, start=block.start, depth=0)
        elif isinstance(block, ir.Table):
            self._convert_table(block.headers, block.rows, block.alignments)
        elif isinstance(block, ir.CodeBlock):
            self._convert_code_block(block.code)
        elif isinstance(block, ir.Image):
            self._convert_image(block.alt, block.path)
        elif isinstance(block, ir.BlockQuote):
            for child in block.children:
                self.convert_block(child)
        elif isinstance(block, ir.ThematicBreak):
            # 水平線 → 空段落で代替
            self.document.add_paragraph()

    # ---- 見出し ----

    def _convert_heading(self, level: int, content: list[ir.Inline]) -> None:
        # 採番カウンタを進める (戻り値は使わないが番号同期のため必要)
        self.heading_mgr.next_heading(level, content)
        if level == 1:
            self.chapter_number = self.heading_mgr.current_h1_number()
            self.figure_in_chapter = 0
            self.table_in_chapter = 0

        plain_text = "".join(ir.inline_to_plain_text(c) for c in content)
        display_text = self.heading_mgr.strip_number(level, plain_text.strip())

        para = self.document.add_paragraph()
        para.style = self.document.styles[str(level)] if str(level) in self.document.styles else None
        # 段落に numbering と keepNext を設定
        pPr = para._p.get_or_add_pPr()
        _set_numbering(pPr, num_id=styles.HEADING_NUM_ID, ilvl=max(level - 1, 0))
        _set_keep_next(pPr)
        # テキスト Run はスタイルが見出しを担うのでフォント・サイズは指定しない
        run = para.add_run(display_text)
        # ただし bold は持たせない (スタイル側で bold)。Run は素のテキストのみ。
        del run

    def _convert_page_break(self) -> None:
        para = self.document.add_paragraph()
        run = para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)

    # ---- 段落 / インライン ----

    def _convert_paragraph(self, content: list[ir.Inline]) -> None:
        para = self.document.add_paragraph()
        if styles.BODY_TEXT_STYLE_ID in self.document.styles:
            para.style = self.document.styles[styles.BODY_TEXT_STYLE_ID]
        self._apply_inlines(para, content, style=_InlineStyle.BODY)

    def _apply_inlines(
        self,
        para,
        inlines: Iterable[ir.Inline],
        *,
        style: _InlineStyle,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        for inline in inlines:
            self._add_inline(para, inline, bold=bold, italic=italic, style=style)

    def _add_inline(
        self,
        para,
        inline: ir.Inline,
        *,
        bold: bool,
        italic: bool,
        style: _InlineStyle,
    ) -> None:
        if isinstance(inline, ir.Text):
            text = _process_text(inline.value)
            self._add_run(para, text, bold=bold, italic=italic, style=style)
        elif isinstance(inline, ir.Code):
            self._add_run(para, f"「{inline.value}」", bold=bold, italic=italic, style=style)
        elif isinstance(inline, ir.Bold):
            for child in inline.children:
                self._add_inline(para, child, bold=True, italic=italic, style=style)
        elif isinstance(inline, ir.Italic):
            for child in inline.children:
                self._add_inline(para, child, bold=bold, italic=True, style=style)
        elif isinstance(inline, ir.Link):
            label_inlines = inline.text
            label_plain = ir.inlines_to_plain_text(label_inlines)
            display = label_plain if label_plain else inline.url
            display = _process_text(display)
            self._add_hyperlink(para, display, inline.url, bold=bold, italic=italic, style=style)
        elif isinstance(inline, ir.SoftBreak):
            self._add_run(para, " ", bold=bold, italic=italic, style=style)
        elif isinstance(inline, ir.HardBreak):
            run = para.add_run()
            br = OxmlElement("w:br")
            run._r.append(br)

    def _add_run(
        self,
        para,
        text: str,
        *,
        bold: bool,
        italic: bool = False,
        style: _InlineStyle,
    ) -> None:
        spec = self._font_spec_for(style)
        run = para.add_run(text)
        run.font.size = None
        # rPr に rFonts と sz を直接書く (フォントの東アジア指定のため)
        rPr = run._r.get_or_add_rPr()
        _set_run_fonts(rPr, ascii_font=spec.ascii_font, east_asia=spec.east_asia)
        _set_run_size(rPr, styles.pt_to_half_point(spec.size_pt))
        if bold:
            _set_run_bold(rPr)
        if italic:
            _set_run_italic(rPr)

    def _font_spec_for(self, style: _InlineStyle) -> _RunFontSpec:
        c = self.config
        if style == _InlineStyle.BODY:
            return _RunFontSpec(c.fonts.body_en, c.fonts.body_ja, c.sizes.body)
        if style == _InlineStyle.TABLE_BODY:
            return _RunFontSpec(c.fonts.body_en, c.fonts.body_ja, c.sizes.table_body)
        return _RunFontSpec(c.fonts.heading_en, c.fonts.heading_ja, c.sizes.table_header)

    def _add_hyperlink(
        self,
        para,
        text: str,
        url: str,
        *,
        bold: bool,
        italic: bool = False,
        style: _InlineStyle,
    ) -> None:
        spec = self._font_spec_for(style)
        hyperlink = OxmlElement("w:hyperlink")
        if url.startswith("#"):
            hyperlink.set(qn("w:anchor"), url[1:])
        else:
            r_id = self.document.part.relate_to(
                url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
            )
            hyperlink.set(qn("r:id"), r_id)

        run_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        _set_run_fonts(rPr, ascii_font=spec.ascii_font, east_asia=spec.east_asia)
        _set_run_size(rPr, styles.pt_to_half_point(spec.size_pt))
        if bold:
            _set_run_bold(rPr)
        if italic:
            _set_run_italic(rPr)
        run_el.append(rPr)
        text_el = OxmlElement("w:t")
        text_el.text = text
        text_el.set(qn("xml:space"), "preserve")
        run_el.append(text_el)
        hyperlink.append(run_el)
        para._p.append(hyperlink)

    # ---- 図表番号 ----

    def _next_figure_number(self) -> str:
        self.figure_seq += 1
        self.figure_in_chapter += 1
        if self.config.numbering.figure_format == "chapter":
            ch = self.chapter_number if self.chapter_number > 0 else 1
            return f"{ch}.{self.figure_in_chapter}"
        return f"{self.figure_seq}"

    def _next_table_number(self) -> str:
        self.table_seq += 1
        self.table_in_chapter += 1
        if self.config.numbering.table_format == "chapter":
            ch = self.chapter_number if self.chapter_number > 0 else 1
            return f"{ch}.{self.table_in_chapter}"
        return f"{self.table_seq}"

    # ---- リスト ----

    def _convert_bullet_list(
        self, items: list[ir.ListItem], depth: int
    ) -> None:
        for item in items:
            level = min(depth, 2)
            para = self.document.add_paragraph()
            if styles.BULLET_STYLE_ID in self.document.styles:
                para.style = self.document.styles[styles.BULLET_STYLE_ID]
            pPr = para._p.get_or_add_pPr()
            _set_numbering(pPr, num_id=styles.BULLET_NUM_ID, ilvl=level)
            self._apply_inlines(para, item.content, bold=False, style=_InlineStyle.BODY)
            for child in item.children:
                if isinstance(child, ir.BulletList):
                    self._convert_bullet_list(child.items, depth + 1)
                elif isinstance(child, ir.OrderedList):
                    self._convert_ordered_list(child.items, child.start, depth + 1)
                else:
                    self.convert_block(child)

    def _convert_ordered_list(
        self, items: list[ir.ListItem], start: int, depth: int
    ) -> None:
        for i, item in enumerate(items):
            num = start + i
            indent_twip = (depth + 1) * styles.pt_to_twip(18.0)
            para = self.document.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            _set_paragraph_indent(pPr, left=indent_twip)
            self._add_run(para, f"{num}. ", bold=False, style=_InlineStyle.BODY)
            self._apply_inlines(para, item.content, bold=False, style=_InlineStyle.BODY)
            for child in item.children:
                if isinstance(child, ir.OrderedList):
                    self._convert_ordered_list(child.items, child.start, depth + 1)
                elif isinstance(child, ir.BulletList):
                    self._convert_bullet_list(child.items, depth + 1)
                else:
                    self.convert_block(child)

    # ---- 表 ----

    def _convert_table(
        self,
        headers: list[list[ir.Inline]],
        rows: list[list[list[ir.Inline]]],
        alignments: list[ir.Alignment],
    ) -> None:
        c = self.config
        body_size_half = styles.pt_to_half_point(c.sizes.body)

        # 表番号キャプション
        table_number = self._next_table_number()
        caption = self.document.add_paragraph()
        _set_paragraph_alignment(caption._p.get_or_add_pPr(), "center")
        if c.numbering.table_format == "chapter":
            self._add_caption_run(
                caption,
                f"表{table_number}",
                ascii_font=c.fonts.heading_en,
                east_asia=c.fonts.heading_ja,
                size_half=body_size_half,
                bold=True,
            )
        else:
            self._add_caption_run(
                caption,
                "表",
                ascii_font=c.fonts.heading_en,
                east_asia=c.fonts.heading_ja,
                size_half=body_size_half,
                bold=True,
            )
            self._add_seq_field(
                caption,
                instr=" SEQ Table \\* ARABIC ",
                visible_text=table_number,
                ascii_font=c.fonts.heading_en,
                east_asia=c.fonts.heading_ja,
                size_half=body_size_half,
                bold=True,
            )

        column_count = max(len(headers), max((len(row) for row in rows), default=0))
        if column_count == 0:
            return
        column_widths = _build_table_grid(column_count, c.page)

        table = self.document.add_table(rows=1 + len(rows), cols=column_count)
        # 表全体の見た目: width 5000 pct, 中央寄せ, fixed layout, セル余白 80 twip
        _apply_table_global_props(table, column_widths)

        # ヘッダー行
        header_row = table.rows[0]
        for col_idx in range(column_count):
            cell = header_row.cells[col_idx]
            _set_cell_width(cell, column_widths[col_idx])
            _set_cell_vertical_align(cell, "center")
            cell.text = ""
            cell_para = cell.paragraphs[0]
            _set_paragraph_alignment(cell_para._p.get_or_add_pPr(), "center")
            if col_idx < len(headers):
                self._apply_inlines(
                    cell_para,
                    headers[col_idx],
                    bold=True,
                    style=_InlineStyle.TABLE_HEADER,
                )

        # データ行
        for r_idx, row in enumerate(rows):
            doc_row = table.rows[r_idx + 1]
            for col_idx in range(column_count):
                cell = doc_row.cells[col_idx]
                _set_cell_width(cell, column_widths[col_idx])
                _set_cell_vertical_align(cell, "center")
                cell.text = ""
                cell_para = cell.paragraphs[0]
                align = _alignment_to_str(alignments[col_idx]) if col_idx < len(alignments) else "left"
                _set_paragraph_alignment(cell_para._p.get_or_add_pPr(), align)
                if col_idx < len(row):
                    self._apply_inlines(
                        cell_para,
                        row[col_idx],
                        bold=False,
                        style=_InlineStyle.TABLE_BODY,
                    )

    # ---- コードブロック ----

    def _convert_code_block(self, code: str) -> None:
        for line in code.splitlines() or [""]:
            para = self.document.add_paragraph()
            run = para.add_run(line)
            rPr = run._r.get_or_add_rPr()
            _set_run_fonts(
                rPr,
                ascii_font="Courier New",
                east_asia="ＭＳ ゴシック",
            )
            _set_run_size(rPr, styles.pt_to_half_point(9.0))

    # ---- 画像 ----

    def _convert_image(self, alt: str, path: str) -> None:
        c = self.config
        body_size_half = styles.pt_to_half_point(c.sizes.body)
        image_path = Path(self.base_path) / path
        try:
            buf = image_path.read_bytes()
        except OSError as e:
            print(
                f"警告: 画像ファイルを読み込めません: {image_path} ({e})",
                file=__import__("sys").stderr,
            )
            self._add_image_fallback(alt)
            return
        try:
            png_buf, width_px, height_px = _convert_to_png_with_dimensions(buf)
        except Exception as e:  # noqa: BLE001
            print(
                f"警告: 画像の変換に失敗しました: {path} ({e})",
                file=__import__("sys").stderr,
            )
            self._add_image_fallback(alt)
            return

        width_emu, height_emu = _fit_image_to_body_width(width_px, height_px, c.page)

        image_para = self.document.add_paragraph()
        _set_paragraph_alignment(image_para._p.get_or_add_pPr(), "center")
        run = image_para.add_run()
        run.add_picture(io.BytesIO(png_buf), width=Emu(width_emu), height=Emu(height_emu))

        # 図番号キャプション
        figure_number = self._next_figure_number()
        caption = self.document.add_paragraph()
        _set_paragraph_alignment(caption._p.get_or_add_pPr(), "center")
        if c.numbering.figure_format == "chapter":
            self._add_caption_run(
                caption,
                f"図{figure_number}",
                ascii_font=c.fonts.body_en,
                east_asia=c.fonts.body_ja,
                size_half=body_size_half,
                bold=False,
            )
        else:
            self._add_caption_run(
                caption,
                "図",
                ascii_font=c.fonts.body_en,
                east_asia=c.fonts.body_ja,
                size_half=body_size_half,
                bold=False,
            )
            self._add_seq_field(
                caption,
                instr=" SEQ Figure \\* ARABIC ",
                visible_text=figure_number,
                ascii_font=c.fonts.body_en,
                east_asia=c.fonts.body_ja,
                size_half=body_size_half,
                bold=False,
            )
        if alt:
            self._add_caption_run(
                caption,
                f" {alt}",
                ascii_font=c.fonts.body_en,
                east_asia=c.fonts.body_ja,
                size_half=body_size_half,
                bold=False,
            )

    def _add_image_fallback(self, alt: str) -> None:
        para = self.document.add_paragraph()
        self._add_run(para, f"[画像: {alt}]", bold=False, style=_InlineStyle.BODY)

    # ---- 共通: キャプション要素 ----

    def _add_caption_run(
        self,
        para,
        text: str,
        *,
        ascii_font: str,
        east_asia: str,
        size_half: int,
        bold: bool,
    ) -> None:
        run = para.add_run(text)
        rPr = run._r.get_or_add_rPr()
        _set_run_fonts(rPr, ascii_font=ascii_font, east_asia=east_asia)
        _set_run_size(rPr, size_half)
        if bold:
            _set_run_bold(rPr)

    def _add_seq_field(
        self,
        para,
        *,
        instr: str,
        visible_text: str,
        ascii_font: str,
        east_asia: str,
        size_half: int,
        bold: bool,
    ) -> None:
        """Word の SEQ フィールドを (begin, instr, separate, text, end) の 5 run で書き込む。"""
        rPr_template = OxmlElement("w:rPr")
        _set_run_fonts(rPr_template, ascii_font=ascii_font, east_asia=east_asia)
        _set_run_size(rPr_template, size_half)
        if bold:
            _set_run_bold(rPr_template)

        def make_run(content: etree._Element) -> etree._Element:
            r = OxmlElement("w:r")
            r.append(_clone_rPr(rPr_template))
            r.append(content)
            return r

        def make_fld_char(kind: str) -> etree._Element:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            return fld

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr

        text_el = OxmlElement("w:t")
        text_el.set(qn("xml:space"), "preserve")
        text_el.text = visible_text

        para._p.append(make_run(make_fld_char("begin")))
        para._p.append(make_run(instr_text))
        para._p.append(make_run(make_fld_char("separate")))
        para._p.append(make_run(text_el))
        para._p.append(make_run(make_fld_char("end")))


# --------- ヘルパ群 (モジュールスコープ) ---------

def _clone_rPr(rPr: etree._Element) -> etree._Element:
    return etree.fromstring(etree.tostring(rPr))


def _set_run_fonts(rPr: etree._Element, *, ascii_font: str, east_asia: str) -> None:
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:cs"), ascii_font)
    rPr.insert(0, fonts)


def _set_run_size(rPr: etree._Element, half_pt: int) -> None:
    existing = rPr.find(qn("w:sz"))
    if existing is not None:
        rPr.remove(existing)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(half_pt))
    rPr.append(sz)


def _set_run_bold(rPr: etree._Element) -> None:
    if rPr.find(qn("w:b")) is None:
        rPr.append(OxmlElement("w:b"))


def _set_run_italic(rPr: etree._Element) -> None:
    if rPr.find(qn("w:i")) is None:
        rPr.append(OxmlElement("w:i"))


def _set_numbering(pPr: etree._Element, *, num_id: int, ilvl: int) -> None:
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(num_id_el)
    pPr.append(numPr)


def _set_keep_next(pPr: etree._Element) -> None:
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(OxmlElement("w:keepNext"))


def _set_paragraph_alignment(pPr: etree._Element, value: str) -> None:
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), value)


def _set_paragraph_indent(pPr: etree._Element, *, left: int) -> None:
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"), str(left))


def _alignment_to_str(alignment: ir.Alignment) -> str:
    if alignment == ir.Alignment.CENTER:
        return "center"
    if alignment == ir.Alignment.RIGHT:
        return "right"
    return "left"


def _set_cell_width(cell, width: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width))
    tcW.set(qn("w:type"), "dxa")


def _set_cell_vertical_align(cell, value: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), value)


def _apply_table_global_props(table, column_widths: list[int]) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # 既存子要素を整理して設定し直す
    for tag in ("tblW", "jc", "tblLayout", "tblCellMar"):
        existing = tblPr.find(qn(f"w:{tag}"))
        if existing is not None:
            tblPr.remove(existing)

    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(TABLE_WIDTH_PCT))
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)

    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    cellMar = OxmlElement("w:tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(TABLE_CELL_PADDING_TWIP))
        m.set(qn("w:type"), "dxa")
        cellMar.append(m)
    tblPr.append(cellMar)

    # tblGrid を上書き
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    tblGrid = OxmlElement("w:tblGrid")
    for w in column_widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        gc.set(qn("w:type"), "dxa")
        tblGrid.append(gc)
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)


# --------- 画像処理 ---------

def _convert_to_png_with_dimensions(buf: bytes) -> tuple[bytes, int, int]:
    img = PILImage.open(io.BytesIO(buf))
    img.load()
    width, height = img.size
    out = io.BytesIO()
    rgba = img if img.mode in ("RGB", "RGBA", "L", "LA", "P") else img.convert("RGBA")
    rgba.save(out, format="PNG")
    return out.getvalue(), width, height


def _fit_image_to_body_width(
    width_px: int, height_px: int, page: PageConfig
) -> tuple[int, int]:
    width_emu = width_px * EMU_PER_PIXEL
    height_emu = height_px * EMU_PER_PIXEL
    body_twip = max(
        page.width - max(page.margin_left, 0) - max(page.margin_right, 0), 0
    )
    max_width_emu = body_twip * EMU_PER_TWIP
    if width_emu <= max_width_emu:
        return int(width_emu), int(height_emu)
    scaled_height = height_emu * max_width_emu // width_emu
    return int(max_width_emu), int(scaled_height)


def _body_width_twip(page: PageConfig) -> int:
    return max(page.width - max(page.margin_left, 0) - max(page.margin_right, 0), 0)


def _build_table_grid(column_count: int, page: PageConfig) -> list[int]:
    body_width = max(_body_width_twip(page), column_count)
    base = body_width // column_count
    remainder = body_width % column_count
    return [base + (1 if i < remainder else 0) for i in range(column_count)]


# --------- 英日間スペース削除 ---------

def _process_text(text: str) -> str:
    if not text:
        return text
    chars = list(text)
    result: list[str] = []
    i = 0
    n = len(chars)
    while i < n:
        c = chars[i]
        if c == " " and 0 < i < n - 1:
            prev = chars[i - 1]
            nxt = chars[i + 1]
            if (_is_ascii_char(prev) and _is_japanese_char(nxt)) or (
                _is_japanese_char(prev) and _is_ascii_char(nxt)
            ):
                i += 1
                continue
        result.append(c)
        i += 1
    return "".join(result)


def _is_ascii_char(c: str) -> bool:
    if not c:
        return False
    return c.isascii() and (c.isalnum() or _is_ascii_punctuation(c))


def _is_ascii_punctuation(c: str) -> bool:
    # Rust の char::is_ascii_punctuation 相当
    code = ord(c)
    return (
        0x21 <= code <= 0x2F
        or 0x3A <= code <= 0x40
        or 0x5B <= code <= 0x60
        or 0x7B <= code <= 0x7E
    )


def _is_japanese_char(c: str) -> bool:
    if not c:
        return False
    code = ord(c)
    return (
        0x3040 <= code <= 0x309F  # ひらがな
        or 0x30A0 <= code <= 0x30FF  # カタカナ
        or 0x4E00 <= code <= 0x9FFF  # CJK 統合漢字
        or 0x3400 <= code <= 0x4DBF  # CJK 統合漢字拡張 A
        or 0xFF00 <= code <= 0xFFEF  # 全角文字
        or 0x3000 <= code <= 0x303F  # CJK 記号
    )
